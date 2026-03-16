# analyzer/services/resume_exporter.py

import io
import re
import logging

logger = logging.getLogger(__name__)


# ─────────────────────────────────────────────────────
# PDF EXPORTER
# ─────────────────────────────────────────────────────
class PDFExporter:
    """
    Converts resume text into a styled PDF.

    Uses ReportLab to generate a clean, ATS-friendly
    single-page (or multi-page) PDF with proper
    typography and section formatting.

    Usage:
        exporter = PDFExporter(resume_text, candidate_name)
        pdf_bytes = exporter.generate()
    """

    # ── Colors (RGB 0-1) ──────────────────────────────
    COLOR_PRIMARY   = (0.10, 0.10, 0.20)   # Near black
    COLOR_ACCENT    = (0.25, 0.31, 0.92)   # Blue
    COLOR_SECTION   = (0.08, 0.08, 0.16)   # Dark header
    COLOR_TEXT      = (0.20, 0.20, 0.28)   # Body text
    COLOR_MUTED     = (0.45, 0.50, 0.58)   # Light text
    COLOR_LINE      = (0.88, 0.88, 0.92)   # Divider

    # ── Page margins ──────────────────────────────────
    MARGIN_LEFT   = 48
    MARGIN_RIGHT  = 48
    MARGIN_TOP    = 50
    MARGIN_BOTTOM = 45

    def __init__(self, resume_text, candidate_name=''):
        self.resume_text    = resume_text
        self.candidate_name = candidate_name
        self.lines          = resume_text.strip().split('\n')

    def generate(self):
        """
        Returns PDF as bytes (ready to stream to browser).
        """
        from reportlab.lib.pagesizes import A4
        from reportlab.pdfgen import canvas
        from reportlab.lib.units import mm
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont

        buf  = io.BytesIO()
        w, h = A4   # 595 x 842 points

        c = canvas.Canvas(buf, pagesize=A4)
        c.setTitle(
            f"Resume — {self.candidate_name}"
            if self.candidate_name else "Optimized Resume"
        )

        # Track vertical position
        y = h - self.MARGIN_TOP

        # ── Process each line ──────────────────────────
        for raw_line in self.lines:
            line = raw_line.rstrip()

            # New page if running out of space
            if y < self.MARGIN_BOTTOM + 30:
                c.showPage()
                y = h - self.MARGIN_TOP

            # ── Detect line type and render ────────────
            if self._is_name_line(line, y, h):
                y = self._draw_name(c, line, w, y)

            elif self._is_contact_line(line):
                y = self._draw_contact(c, line, w, y)

            elif self._is_section_header(line):
                y = self._draw_section_header(c, line, w, y)

            elif self._is_bullet(line):
                y = self._draw_bullet(c, line, w, y)

            elif line.strip() == '':
                y -= 5   # Small gap for blank lines

            else:
                y = self._draw_body(c, line, w, y)

        c.save()
        buf.seek(0)
        return buf.read()

    # ─────────────────────────────────────────────────
    # LINE TYPE DETECTION
    # ─────────────────────────────────────────────────
    def _is_name_line(self, line, y, h):
        """First non-empty line at top = candidate name."""
        return (
            y > h * 0.80 and
            line.strip() and
            len(line.strip()) < 60 and
            not self._is_contact_line(line) and
            not '@' in line
        )

    def _is_contact_line(self, line):
        """Lines with email, phone, linkedin, github."""
        lower = line.lower()
        return any(kw in lower for kw in [
            '@', 'linkedin', 'github', 'phone',
            '+91', '+1', 'http', 'www.',
        ]) or bool(re.search(r'\d{5,}', line))

    def _is_section_header(self, line):
        """
        All-caps short line, or known section keywords.
        """
        stripped = line.strip()
        if not stripped or len(stripped) > 60:
            return False

        known = [
            'SUMMARY', 'OBJECTIVE', 'PROFILE',
            'EXPERIENCE', 'INTERNSHIP', 'WORK',
            'EDUCATION', 'SKILLS', 'TECHNICAL',
            'PROJECTS', 'CERTIFICATIONS', 'ACHIEVEMENTS',
            'LANGUAGES', 'INTERESTS', 'REFERENCES',
        ]
        upper = stripped.upper()
        return (
            stripped.isupper() or
            any(kw in upper for kw in known)
        )

    def _is_bullet(self, line):
        """Lines starting with -, •, *, ●."""
        return line.strip().startswith(
            ('-', '•', '*', '●', '▪')
        )

    # ─────────────────────────────────────────────────
    # DRAWING METHODS
    # ─────────────────────────────────────────────────
    def _draw_name(self, c, line, w, y):
        """Draws the candidate name — large, centered."""
        c.setFont('Helvetica-Bold', 20)
        c.setFillColorRGB(*self.COLOR_PRIMARY)
        c.drawCentredString(w / 2, y, line.strip())

        # Accent underline
        name_w = c.stringWidth(line.strip(),
                                'Helvetica-Bold', 20)
        cx = w / 2
        c.setStrokeColorRGB(*self.COLOR_ACCENT)
        c.setLineWidth(2)
        c.line(cx - name_w/2, y - 4,
               cx + name_w/2, y - 4)

        return y - 28

    def _draw_contact(self, c, line, w, y):
        """Draws contact info — small, centered, muted."""
        c.setFont('Helvetica', 8.5)
        c.setFillColorRGB(*self.COLOR_MUTED)

        # Wrap if too long
        max_w  = w - self.MARGIN_LEFT - self.MARGIN_RIGHT
        text   = line.strip()

        if c.stringWidth(text, 'Helvetica', 8.5) > max_w:
            # Split at | or ,
            parts = re.split(r'[|,]', text)
            mid   = len(parts) // 2
            line1 = ' | '.join(p.strip() for p in parts[:mid])
            line2 = ' | '.join(p.strip() for p in parts[mid:])
            c.drawCentredString(w/2, y, line1)
            y -= 13
            c.drawCentredString(w/2, y, line2)
        else:
            c.drawCentredString(w/2, y, text)

        return y - 14

    def _draw_section_header(self, c, line, w, y):
        """Draws section header with full-width underline."""
        y -= 8  # Space before section

        # Header text
        c.setFont('Helvetica-Bold', 11)
        c.setFillColorRGB(*self.COLOR_PRIMARY)
        c.drawString(self.MARGIN_LEFT, y, line.strip().upper())

        # Full-width line
        y -= 4
        c.setStrokeColorRGB(*self.COLOR_LINE)
        c.setLineWidth(0.8)
        c.line(self.MARGIN_LEFT, y,
               w - self.MARGIN_RIGHT, y)

        return y - 12

    def _draw_bullet(self, c, line, w, y):
        """Draws a bullet point with hanging indent."""
        text      = line.strip().lstrip('-•*● ▪').strip()
        indent    = self.MARGIN_LEFT + 14
        bullet_x  = self.MARGIN_LEFT + 4
        max_w     = w - indent - self.MARGIN_RIGHT

        c.setFont('Helvetica', 9.5)
        c.setFillColorRGB(*self.COLOR_ACCENT)
        c.drawString(bullet_x, y, '•')

        c.setFillColorRGB(*self.COLOR_TEXT)
        wrapped = self._wrap_text(
            text, 'Helvetica', 9.5, max_w
        )

        for i, wline in enumerate(wrapped):
            c.drawString(indent, y, wline)
            if i < len(wrapped) - 1:
                y -= 13

        return y - 15

    def _draw_body(self, c, line, w, y):
        """Draws regular body text."""
        text  = line.strip()
        max_w = w - self.MARGIN_LEFT - self.MARGIN_RIGHT

        # Job title / company line detection
        is_job = (
            bool(re.search(r'(20\d{2}|present|current)',
                           text, re.I)) or
            '—' in text or ' - ' in text
        )

        if is_job:
            c.setFont('Helvetica-Bold', 9.5)
            c.setFillColorRGB(*self.COLOR_PRIMARY)
        else:
            c.setFont('Helvetica', 9.5)
            c.setFillColorRGB(*self.COLOR_TEXT)

        wrapped = self._wrap_text(
            text,
            'Helvetica-Bold' if is_job else 'Helvetica',
            9.5,
            max_w
        )

        for i, wline in enumerate(wrapped):
            c.drawString(self.MARGIN_LEFT, y, wline)
            if i < len(wrapped) - 1:
                y -= 13

        return y - 14

    # ─────────────────────────────────────────────────
    # TEXT WRAPPING
    # ─────────────────────────────────────────────────
    def _wrap_text(self, text, font, size, max_width):
        """
        Wraps text to fit within max_width.
        Returns list of lines.
        """
        from reportlab.pdfgen import canvas as _c
        from reportlab.lib.pagesizes import A4

        # Temp canvas just for measuring
        tmp   = io.BytesIO()
        tmp_c = _c.Canvas(tmp, pagesize=A4)

        words  = text.split()
        lines  = []
        current = ''

        for word in words:
            test = (current + ' ' + word).strip()
            if tmp_c.stringWidth(test, font, size) <= max_width:
                current = test
            else:
                if current:
                    lines.append(current)
                current = word

        if current:
            lines.append(current)

        return lines or ['']


# ─────────────────────────────────────────────────────
# DOCX EXPORTER
# ─────────────────────────────────────────────────────
class DOCXExporter:
    """
    Converts resume text into a formatted .docx file.

    Uses python-docx to create a clean Word document
    with proper styles, sections, and formatting.

    Usage:
        exporter = DOCXExporter(resume_text, candidate_name)
        docx_bytes = exporter.generate()
    """

    def __init__(self, resume_text, candidate_name=''):
        self.resume_text    = resume_text
        self.candidate_name = candidate_name
        self.lines          = resume_text.strip().split('\n')

    def generate(self):
        """
        Returns DOCX as bytes.
        """
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        doc = Document()

        # ── Page margins ──────────────────────────────
        for section in doc.sections:
            section.top_margin    = Cm(1.5)
            section.bottom_margin = Cm(1.5)
            section.left_margin   = Cm(1.8)
            section.right_margin  = Cm(1.8)

        # ── Process lines ──────────────────────────────
        first_line    = True
        contact_block = []

        for raw_line in self.lines:
            line = raw_line.rstrip()

            # ── Name (first non-empty line) ────────────
            if first_line and line.strip():
                self._add_name(doc, line.strip())
                first_line = False
                continue

            # ── Contact lines (batch them) ─────────────
            if self._is_contact_line(line):
                contact_block.append(line.strip())
                continue
            elif contact_block:
                self._add_contact(
                    doc, ' | '.join(contact_block)
                )
                contact_block = []

            # ── Section header ─────────────────────────
            if self._is_section_header(line):
                self._add_section_header(doc, line.strip())
                continue

            # ── Bullet ────────────────────────────────
            if self._is_bullet(line):
                text = line.strip().lstrip('-•*● ▪').strip()
                self._add_bullet(doc, text)
                continue

            # ── Blank line ─────────────────────────────
            if not line.strip():
                continue

            # ── Body text ──────────────────────────────
            self._add_body(doc, line.strip())

        # Flush any remaining contact lines
        if contact_block:
            self._add_contact(
                doc, ' | '.join(contact_block)
            )

        # ── Save to bytes ──────────────────────────────
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf.read()

    # ─────────────────────────────────────────────────
    # LINE TYPE HELPERS
    # ─────────────────────────────────────────────────
    def _is_contact_line(self, line):
        lower = line.lower()
        return any(kw in lower for kw in [
            '@', 'linkedin', 'github', '+91', '+1',
            'phone', 'http', 'www.',
        ]) or bool(re.search(r'\d{7,}', line))

    def _is_section_header(self, line):
        stripped = line.strip()
        if not stripped or len(stripped) > 65:
            return False
        known = [
            'SUMMARY', 'OBJECTIVE', 'EXPERIENCE',
            'INTERNSHIP', 'EDUCATION', 'SKILLS',
            'PROJECTS', 'CERTIFICATIONS', 'ACHIEVEMENTS',
            'TECHNICAL', 'PROFILE', 'WORK',
        ]
        upper = stripped.upper()
        return (
            stripped.isupper() or
            any(kw in upper for kw in known)
        )

    def _is_bullet(self, line):
        return line.strip().startswith(
            ('-', '•', '*', '●', '▪')
        )

    # ─────────────────────────────────────────────────
    # DOCX DRAWING METHODS
    # ─────────────────────────────────────────────────
    def _add_name(self, doc, text):
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        p    = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run  = p.add_run(text)
        run.bold      = True
        run.font.size = Pt(22)
        run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

        # Space after
        p.paragraph_format.space_after = Pt(4)

    def _add_contact(self, doc, text):
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        p    = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run  = p.add_run(text)
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x64, 0x74, 0x8b)

        p.paragraph_format.space_after = Pt(2)

    def _add_section_header(self, doc, text):
        from docx.shared import Pt, RGBColor
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        p    = doc.add_paragraph()
        run  = p.add_run(text.upper())
        run.bold      = True
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

        # Bottom border (section divider)
        pPr  = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'),   'single')
        bottom.set(qn('w:sz'),    '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), 'ccccdd')
        pBdr.append(bottom)
        pPr.append(pBdr)

        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after  = Pt(4)

    def _add_bullet(self, doc, text):
        from docx.shared import Pt, RGBColor, Inches

        p   = doc.add_paragraph(style='List Bullet')
        run = p.add_run(text)
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(0x33, 0x3a, 0x4a)

        p.paragraph_format.space_after   = Pt(2)
        p.paragraph_format.left_indent   = Inches(0.25)

    def _add_body(self, doc, text):
        from docx.shared import Pt, RGBColor

        # Detect job title / company line
        is_job = (
            bool(re.search(
                r'(20\d{2}|present|current)',
                text, re.I
            )) or '—' in text or ' - ' in text
        )

        p   = doc.add_paragraph()
        run = p.add_run(text)
        run.font.size = Pt(9.5)

        if is_job:
            run.bold = True
            run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
        else:
            run.font.color.rgb = RGBColor(0x33, 0x3a, 0x4a)

        p.paragraph_format.space_after = Pt(2)


# ─────────────────────────────────────────────────────
# HELPER FUNCTIONS
# ─────────────────────────────────────────────────────
def export_as_pdf(resume_text, candidate_name=''):
    """Returns PDF bytes."""
    return PDFExporter(resume_text, candidate_name).generate()


def export_as_docx(resume_text, candidate_name=''):
    """Returns DOCX bytes."""
    return DOCXExporter(resume_text, candidate_name).generate()


def extract_candidate_name(resume_text):
    """
    Tries to extract candidate name from first line
    of resume text.
    """
    lines = resume_text.strip().split('\n')
    for line in lines[:3]:
        stripped = line.strip()
        if (stripped and
                len(stripped) < 50 and
                '@' not in stripped and
                not stripped[0].isdigit()):
            return stripped
    return 'Candidate'