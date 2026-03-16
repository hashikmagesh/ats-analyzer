# analyzer/services/resume_extractor.py

import re
import logging
import pdfplumber
import docx
from pathlib import Path

# Logger — helps us debug issues without crashing the app
logger = logging.getLogger(__name__)


class ResumeExtractor:
    """
    Extracts and cleans text from PDF and DOCX resume files.

    Usage:
        extractor = ResumeExtractor(file_path)
        result = extractor.extract()
        print(result['text'])
    """

    # File types we support
    SUPPORTED_FORMATS = ['.pdf', '.docx']

    def __init__(self, file_path):
        """
        file_path: full path to the resume file on disk
        e.g. "C:/Users/.../media/resumes/john_resume.pdf"
        """
        self.file_path = Path(file_path)
        self.extension = self.file_path.suffix.lower()

    # ─────────────────────────────────────────────────
    # MAIN METHOD — call this to extract text
    # ─────────────────────────────────────────────────
    def extract(self):
        """
        Main entry point. Detects file type and routes
        to the correct extraction method.

        Returns a dict:
        {
            'success': True/False,
            'text': 'cleaned resume text...',
            'page_count': 2,
            'word_count': 450,
            'error': None or 'error message'
        }
        """
        # Check file exists
        if not self.file_path.exists():
            return self._error_result(f"File not found: {self.file_path}")

        # Check file format is supported
        if self.extension not in self.SUPPORTED_FORMATS:
            return self._error_result(
                f"Unsupported format: {self.extension}. "
                f"Use PDF or DOCX."
            )

        # Route to the correct extractor
        if self.extension == '.pdf':
            return self._extract_from_pdf()
        elif self.extension == '.docx':
            return self._extract_from_docx()

    # ─────────────────────────────────────────────────
    # PDF EXTRACTION
    # ─────────────────────────────────────────────────
    def _extract_from_pdf(self):
        """
        Extracts text from PDF using pdfplumber.
        pdfplumber is better than PyPDF2 because it:
        - Handles multi-column layouts
        - Preserves text order better
        - Handles tables inside PDFs
        """
        try:
            all_text = []
            page_count = 0

            with pdfplumber.open(self.file_path) as pdf:
                page_count = len(pdf.pages)

                for page_num, page in enumerate(pdf.pages, start=1):
                    # Extract text from this page
                    page_text = page.extract_text()

                    if page_text:
                        all_text.append(page_text)
                        logger.debug(f"Page {page_num}: extracted "
                                     f"{len(page_text)} characters")
                    else:
                        logger.warning(f"Page {page_num}: no text found "
                                       f"(might be image-based)")

            # Join all pages together
            raw_text = '\n'.join(all_text)

            # Clean the text
            cleaned_text = self._clean_text(raw_text)

            if not cleaned_text:
                return self._error_result(
                    "No text could be extracted. "
                    "The PDF might be image-based (scanned). "
                    "Please use a text-based PDF."
                )

            return self._success_result(cleaned_text, page_count)

        except Exception as e:
            logger.error(f"PDF extraction failed: {str(e)}")
            return self._error_result(f"Failed to read PDF: {str(e)}")

    # ─────────────────────────────────────────────────
    # DOCX EXTRACTION
    # ─────────────────────────────────────────────────
    def _extract_from_docx(self):
        """
        Extracts text from DOCX using python-docx.
        Handles paragraphs AND tables inside Word documents.
        """
        try:
            doc = docx.Document(self.file_path)
            all_text = []

            # ── Extract from paragraphs ──────────────
            # Paragraphs are the main text blocks in Word
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:  # Skip empty paragraphs
                    all_text.append(text)

            # ── Extract from tables ──────────────────
            # Many resumes use tables for layout
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text = cell.text.strip()
                        if text:
                            all_text.append(text)

            # Join everything
            raw_text = '\n'.join(all_text)

            # Clean the text
            cleaned_text = self._clean_text(raw_text)

            if not cleaned_text:
                return self._error_result(
                    "No text could be extracted from the DOCX file."
                )

            # DOCX doesn't have pages, estimate from word count
            word_count = len(cleaned_text.split())
            estimated_pages = max(1, word_count // 400)

            return self._success_result(cleaned_text, estimated_pages)

        except Exception as e:
            logger.error(f"DOCX extraction failed: {str(e)}")
            return self._error_result(f"Failed to read DOCX: {str(e)}")

    # ─────────────────────────────────────────────────
    # TEXT CLEANING
    # ─────────────────────────────────────────────────
    def _clean_text(self, text):
        """
        Cleans raw extracted text.

        Problems we fix:
        1. Unicode garbage characters (â€™ instead of ')
        2. Multiple blank lines → single blank line
        3. Multiple spaces → single space
        4. Leading/trailing whitespace
        5. Weird bullet point symbols
        """
        if not text:
            return ""

        # ── Step 1: Fix encoding issues ──────────────
        # Replace common unicode problems
        text = text.encode('utf-8', errors='ignore').decode('utf-8')

        # ── Step 2: Normalize bullet points ──────────
        # Many resumes use fancy bullets (•, ●, ◆, ➢, ▪)
        # Replace them all with a standard dash
        bullet_chars = ['•', '●', '◆', '◇', '➢', '➤', '▪', '▸', '✓', '✔', '→', '►']
        for bullet in bullet_chars:
            text = text.replace(bullet, '-')

        # ── Step 3: Remove special/control characters ─
        # Keep only printable characters + newlines + tabs
        text = re.sub(r'[^\x20-\x7E\n\t]', ' ', text)

        # ── Step 4: Clean up whitespace ───────────────
        # Multiple spaces → single space
        text = re.sub(r'[ \t]+', ' ', text)

        # More than 2 newlines → 2 newlines
        text = re.sub(r'\n{3,}', '\n\n', text)

        # Clean each line
        lines = []
        for line in text.split('\n'):
            line = line.strip()
            if line:  # Skip completely empty lines
                lines.append(line)

        # ── Step 5: Final cleanup ─────────────────────
        text = '\n'.join(lines)
        text = text.strip()

        return text

    # ─────────────────────────────────────────────────
    # HELPER METHODS
    # ─────────────────────────────────────────────────
    def _success_result(self, text, page_count):
        """Returns a standardized success dictionary."""
        word_count = len(text.split())
        char_count = len(text)

        return {
            'success': True,
            'text': text,
            'page_count': page_count,
            'word_count': word_count,
            'char_count': char_count,
            'error': None,
        }

    def _error_result(self, error_message):
        """Returns a standardized error dictionary."""
        return {
            'success': False,
            'text': '',
            'page_count': 0,
            'word_count': 0,
            'char_count': 0,
            'error': error_message,
        }


# ─────────────────────────────────────────────────────
# STANDALONE HELPER FUNCTION
# ─────────────────────────────────────────────────────
def extract_resume_text(file_path):
    """
    Simple helper function — shortcut to use the extractor.

    Usage:
        from analyzer.services.resume_extractor import extract_resume_text
        result = extract_resume_text('/path/to/resume.pdf')
    """
    extractor = ResumeExtractor(file_path)
    return extractor.extract()